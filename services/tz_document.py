"""
TZ Document Generator
=====================
Генерация красивых .docx документов ТЗ.
"""

import os
import logging
from datetime import datetime
from typing import Optional, Any

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn

from config import config

logger = logging.getLogger(__name__)


class TZDocumentGenerator:
    """
    Генератор документов ТЗ в формате .docx
    Создаёт красивые, структурированные документы для клиентов.
    """
    
    def __init__(self):
        self.temp_dir = config.TEMP_DIR
        os.makedirs(self.temp_dir, exist_ok=True)
    
    def generate_tz_docx(
        self,
        brief_data: dict,
        user_id: int = 0,
        include_empty_sections: bool = False
    ) -> str:
        """
        Генерирует ТЗ в формате .docx
        
        Args:
            brief_data: Словарь с данными брифа
            user_id: ID пользователя (для имени файла)
            include_empty_sections: Показывать ли пустые разделы
            
        Returns:
            Путь к сгенерированному файлу
        """
        doc = Document()
        
        # Настраиваем стили
        self._setup_styles(doc)
        
        # === ШАПКА ДОКУМЕНТА ===
        self._add_header(doc, brief_data)
        
        # === 1. ЦЕЛЬ ПРОЕКТА ===
        self._add_section(
            doc, 
            "1. Цель проекта",
            brief_data.get("project_goal"),
            required=True
        )
        
        # === 2. ЦЕЛЕВАЯ АУДИТОРИЯ ===
        self._add_section(
            doc,
            "2. Целевая аудитория",
            brief_data.get("target_audience"),
            include_empty=include_empty_sections
        )
        
        # === 3. ТИП ПРОЕКТА И ПЛАТФОРМА ===
        type_platform_text = self._format_type_platform(brief_data)
        self._add_section(
            doc,
            "3. Тип проекта и платформа",
            type_platform_text,
            required=True
        )
        
        # === 4. ФУНКЦИОНАЛЬНЫЕ ТРЕБОВАНИЯ ===
        self._add_features_section(doc, brief_data)
        
        # === 5. ИНТЕГРАЦИИ ===
        if brief_data.get("integrations"):
            self._add_list_section(
                doc,
                "5. Интеграции",
                brief_data.get("integrations")
            )
        
        # === 6. РЕФЕРЕНСЫ ===
        if brief_data.get("references"):
            self._add_list_section(
                doc,
                "6. Референсы",
                brief_data.get("references")
            )
        
        # === 7. ОГРАНИЧЕНИЯ И УСЛОВИЯ ===
        self._add_constraints_section(doc, brief_data)
        
        # === 8. ОЖИДАЕМЫЕ РЕЗУЛЬТАТЫ ===
        if brief_data.get("deliverables"):
            self._add_list_section(
                doc,
                "8. Ожидаемые результаты (Deliverables)",
                brief_data.get("deliverables")
            )
        elif include_empty_sections:
            self._add_section(
                doc,
                "8. Ожидаемые результаты",
                "⚠️ Требуется уточнение",
                include_empty=True
            )
        
        # === 9. КРИТЕРИИ ПРИЁМКИ ===
        if brief_data.get("acceptance_criteria"):
            self._add_list_section(
                doc,
                "9. Критерии приёмки",
                brief_data.get("acceptance_criteria")
            )
        
        # === 10. РИСКИ И КРАСНЫЕ ФЛАГИ ===
        if brief_data.get("risks"):
            self._add_risks_section(doc, brief_data.get("risks"))
        
        # === 11. ОТКРЫТЫЕ ВОПРОСЫ ===
        if brief_data.get("open_questions"):
            self._add_questions_section(doc, brief_data.get("open_questions"))
        
        # === ПОДВАЛ ===
        self._add_footer(doc)
        
        # Сохраняем файл
        project_name = brief_data.get("project_name", "").replace(" ", "_")[:20]
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        
        if project_name:
            filename = f"TZ_{project_name}_{timestamp}.docx"
        else:
            filename = f"TZ_{user_id}_{timestamp}.docx"
        
        filepath = os.path.join(self.temp_dir, filename)
        doc.save(filepath)
        
        logger.info(f"Сгенерирован документ: {filepath}")
        return filepath
    
    def _setup_styles(self, doc: Document):
        """Настраивает стили документа"""
        # Основной шрифт
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(11)
        
        # Настраиваем стиль заголовков
        for i in range(1, 4):
            style_name = f'Heading {i}'
            if style_name in doc.styles:
                heading_style = doc.styles[style_name]
                heading_style.font.name = 'Arial'
                heading_style.font.bold = True
                if i == 1:
                    heading_style.font.size = Pt(16)
                    heading_style.font.color.rgb = RGBColor(0, 51, 102)
                elif i == 2:
                    heading_style.font.size = Pt(14)
                    heading_style.font.color.rgb = RGBColor(0, 76, 153)
                else:
                    heading_style.font.size = Pt(12)
    
    def _add_header(self, doc: Document, brief_data: dict):
        """Добавляет шапку документа"""
        # Заголовок
        title = doc.add_heading('ТЕХНИЧЕСКОЕ ЗАДАНИЕ', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Название проекта (если есть)
        project_name = brief_data.get("project_name")
        if project_name:
            subtitle = doc.add_paragraph()
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = subtitle.add_run(project_name)
            run.bold = True
            run.font.size = Pt(14)
        
        # Дата
        date_para = doc.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_run = date_para.add_run(f"Дата: {datetime.now().strftime('%d.%m.%Y')}")
        date_run.font.size = Pt(10)
        date_run.font.color.rgb = RGBColor(128, 128, 128)
        
        # Разделитель
        doc.add_paragraph()
    
    def _add_section(
        self, 
        doc: Document, 
        title: str, 
        content: Optional[str],
        required: bool = False,
        include_empty: bool = False
    ):
        """Добавляет раздел с заголовком и текстом"""
        if not content and not include_empty and not required:
            return
        
        doc.add_heading(title, level=1)
        
        if content:
            para = doc.add_paragraph(content)
            para.paragraph_format.space_after = Pt(12)
        elif required:
            para = doc.add_paragraph("⚠️ Информация не предоставлена")
            para.runs[0].font.italic = True
            para.runs[0].font.color.rgb = RGBColor(255, 153, 0)
        elif include_empty:
            para = doc.add_paragraph("Требуется уточнение")
            para.runs[0].font.italic = True
            para.runs[0].font.color.rgb = RGBColor(128, 128, 128)
    
    def _add_list_section(
        self, 
        doc: Document, 
        title: str, 
        items: list[str]
    ):
        """Добавляет раздел со списком"""
        if not items:
            return
        
        doc.add_heading(title, level=1)
        
        for item in items:
            para = doc.add_paragraph(style='List Bullet')
            para.add_run(item)
        
        doc.add_paragraph()  # Отступ после списка
    
    def _format_type_platform(self, brief_data: dict) -> str:
        """Форматирует информацию о типе проекта и платформе"""
        lines = []
        
        project_type = brief_data.get("project_type")
        if project_type:
            lines.append(f"Тип проекта: {project_type}")
        
        platform = brief_data.get("platform")
        if platform:
            lines.append(f"Платформа: {platform}")
        
        return "\n".join(lines) if lines else "Не указано"
    
    def _add_features_section(self, doc: Document, brief_data: dict):
        """Добавляет раздел с функциональными требованиями"""
        must_have = brief_data.get("must_have_features", [])
        nice_to_have = brief_data.get("nice_to_have_features", [])
        
        if not must_have and not nice_to_have:
            return
        
        doc.add_heading("4. Функциональные требования", level=1)
        
        if must_have:
            doc.add_heading("Обязательный функционал (Must Have)", level=2)
            for item in must_have:
                para = doc.add_paragraph(style='List Bullet')
                para.add_run(item)
        
        if nice_to_have:
            doc.add_heading("Желательный функционал (Nice to Have)", level=2)
            for item in nice_to_have:
                para = doc.add_paragraph(style='List Bullet')
                run = para.add_run(item)
                run.font.color.rgb = RGBColor(100, 100, 100)
        
        doc.add_paragraph()
    
    def _add_constraints_section(self, doc: Document, brief_data: dict):
        """Добавляет раздел с ограничениями"""
        deadline = brief_data.get("deadline")
        budget = brief_data.get("budget_range")
        constraints = brief_data.get("constraints", [])
        content_ready = brief_data.get("content_ready")
        stakeholders = brief_data.get("stakeholders")
        
        if not any([deadline, budget, constraints, content_ready, stakeholders]):
            return
        
        doc.add_heading("7. Ограничения и условия", level=1)
        
        # Таблица с основными ограничениями
        table = doc.add_table(rows=0, cols=2)
        table.style = 'Table Grid'
        
        if deadline:
            row = table.add_row()
            row.cells[0].text = "Сроки"
            row.cells[1].text = deadline
            row.cells[0].paragraphs[0].runs[0].bold = True
        
        if budget:
            row = table.add_row()
            row.cells[0].text = "Бюджет"
            row.cells[1].text = budget
            row.cells[0].paragraphs[0].runs[0].bold = True
        
        if content_ready:
            row = table.add_row()
            row.cells[0].text = "Контент"
            row.cells[1].text = content_ready
            row.cells[0].paragraphs[0].runs[0].bold = True
        
        if stakeholders:
            row = table.add_row()
            row.cells[0].text = "Принимает решения"
            row.cells[1].text = stakeholders
            row.cells[0].paragraphs[0].runs[0].bold = True
        
        # Дополнительные ограничения
        if constraints:
            doc.add_paragraph()
            doc.add_heading("Дополнительные ограничения:", level=2)
            for c in constraints:
                para = doc.add_paragraph(style='List Bullet')
                para.add_run(c)
        
        doc.add_paragraph()
    
    def _add_risks_section(self, doc: Document, risks: list[str]):
        """Добавляет раздел с рисками"""
        if not risks:
            return
        
        doc.add_heading("10. Риски и предупреждения", level=1)
        
        for risk in risks:
            para = doc.add_paragraph()
            run = para.add_run("⚠️ ")
            run.font.color.rgb = RGBColor(255, 153, 0)
            para.add_run(risk)
        
        doc.add_paragraph()
    
    def _add_questions_section(self, doc: Document, questions: list[str]):
        """Добавляет раздел с открытыми вопросами"""
        if not questions:
            return
        
        doc.add_heading("11. Открытые вопросы", level=1)
        
        intro = doc.add_paragraph(
            "Следующие вопросы требуют уточнения перед началом работ:"
        )
        intro.runs[0].font.italic = True
        
        for i, q in enumerate(questions, 1):
            para = doc.add_paragraph()
            run = para.add_run(f"{i}. {q}")
        
        doc.add_paragraph()
    
    def _add_footer(self, doc: Document):
        """Добавляет подвал документа"""
        doc.add_paragraph()
        
        # Горизонтальная линия (через пустой параграф с границей)
        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(24)
        
        # Подпись
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("Документ сгенерирован AI Brief Refiner")
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(128, 128, 128)
        run.font.italic = True
        
        timestamp = doc.add_paragraph()
        timestamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ts_run = timestamp.add_run(datetime.now().strftime('%d.%m.%Y %H:%M'))
        ts_run.font.size = Pt(9)
        ts_run.font.color.rgb = RGBColor(160, 160, 160)
    
    def cleanup_old_files(self, max_age_hours: int = 24):
        """Удаляет старые временные файлы"""
        import time
        
        current_time = time.time()
        max_age_seconds = max_age_hours * 3600
        
        for filename in os.listdir(self.temp_dir):
            if filename.endswith('.docx'):
                filepath = os.path.join(self.temp_dir, filename)
                if os.path.isfile(filepath):
                    file_age = current_time - os.path.getmtime(filepath)
                    if file_age > max_age_seconds:
                        os.remove(filepath)
                        logger.debug(f"Удалён старый файл: {filename}")


# Синглтон
_generator: Optional[TZDocumentGenerator] = None


def get_tz_generator() -> TZDocumentGenerator:
    """Возвращает синглтон генератора"""
    global _generator
    if _generator is None:
        _generator = TZDocumentGenerator()
    return _generator

